from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


def dict_to_word(data):
    date = datetime.today().strftime("%d.%m.%Y") # получаю текущую дату по формату 01.01.1970

    doc = Document() # инициализирую объект документа
    section = doc.sections[0] # выбираю основную секцию
    section.left_margin = Cm(2) # выравнивание по ширине
    section.right_margin = Cm(2)

    text_1 = doc.add_paragraph() # добавляю объект параграфа
    text_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT # выравнивание
    text_1.paragraph_format.space_before = Pt(0) # отступ до
    text_1.paragraph_format.space_after = Pt(0) # отступ после
    text_1_run1 = text_1.add_run('ИП Холостенко Максим Олегович') # сам текст
    text_1_run1.font.name = 'Calibri' # оформление ....
    text_1_run1.font.size = Pt(12)
    text_1_run1.font.bold = True
    text_1_run1.font.italic = True

    text_2 = doc.add_paragraph()
    text_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    text_2.paragraph_format.space_before = Pt(0)
    text_2.paragraph_format.space_after = Pt(0)
    text_2_run1 = text_2.add_run('ИНН: 750601450118')
    text_2_run1.font.name = 'Calibri'
    text_2_run1.font.size = Pt(12)
    text_2_run1.font.bold = False
    text_2_run1.font.italic = True

    text_3 = doc.add_paragraph()
    text_3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text_3.paragraph_format.space_before = Pt(18)
    text_3.paragraph_format.space_after = Pt(10)
    text_3_run1 = text_3.add_run(f'Коммерческое предложение от {date}г.')
    text_3_run1.font.name = 'Calibri'
    text_3_run1.font.size = Pt(12)
    text_3_run1.font.bold = True
    text_3_run1.font.italic = True

    # Объявляю параметры таблицы, создаю её как объект
    rows = len(data) + 1
    cols = 5
    table = doc.add_table(rows=rows, cols=cols)
    hdr_cells = table.rows[0].cells

    # Заполняю верхние поля таблицы
    hdr_cells[0].text = '№ п\п'
    hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
    hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[0].width = Cm(1)

    hdr_cells[1].text = 'Наименование товара'
    hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
    hdr_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
    hdr_cells[1].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].allow_autofit = False
    hdr_cells[1].width = Inches(8)

    hdr_cells[2].text = 'Кол-во, шт.'
    hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[2].paragraphs[0].runs[0].font.name = 'Calibri'
    hdr_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
    hdr_cells[2].paragraphs[0].runs[0].font.bold = True
    hdr_cells[2].width = Inches(1)

    hdr_cells[3].text = 'Цена, р.'
    hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[3].paragraphs[0].runs[0].font.name = 'Calibri'
    hdr_cells[3].paragraphs[0].runs[0].font.size = Pt(12)
    hdr_cells[3].paragraphs[0].runs[0].font.bold = True

    hdr_cells[4].text = 'Стоимость, р.'
    hdr_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells[4].paragraphs[0].runs[0].font.name = 'Calibri'
    hdr_cells[4].paragraphs[0].runs[0].font.size = Pt(12)
    hdr_cells[4].paragraphs[0].runs[0].font.bold = True

    for el in data: # заполняю поля таблицы в цикле из словаря
        if el == 'total':
            row_cells = table.rows[len(data)].cells
            row_cells[0].merge(row_cells[1])
            row_cells[0].merge(row_cells[2])
            row_cells[0].merge(row_cells[3])
            row_cells[0].text = data['total']['name']
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[4].text = data['total']['cost']
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[4].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[4].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[4].paragraphs[0].runs[0].font.bold = True
            row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            row_cells = table.rows[el + 1].cells
            row_cells[0].text = str(data[el]['num'])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[0].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            row_cells[1].text = data[el]['name']
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row_cells[1].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[1].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            row_cells[2].text = str(data[el]['qty'])
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[2].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[2].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            row_cells[3].text = data[el]['price']
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[3].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[3].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            row_cells[4].text = data[el]['cost']
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[4].paragraphs[0].runs[0].font.name = 'Calibri'
            row_cells[4].paragraphs[0].runs[0].font.size = Pt(12)
            row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.style = 'Table Grid'
    table.allow_autofit = False

    text_4 = doc.add_paragraph()
    text_4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text_4.paragraph_format.space_before = Pt(18)
    text_4.paragraph_format.space_after = Pt(0)
    text_4_run1 = text_4.add_run(f"Итого: {data['total']['prop']}.")
    text_4_run1.font.name = 'Calibri'
    text_4_run1.font.size = Pt(12)

    text_5 = doc.add_paragraph()
    text_5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text_5.paragraph_format.space_before = Pt(0)
    text_5.paragraph_format.space_after = Pt(0)
    text_5_run1 = text_5.add_run('Предложение действительно в течении 5 дней'
                                 '\nПредложение не является публичной офертой.')
    text_5_run1.font.name = 'Calibri'
    text_5_run1.font.size = Pt(12)

    doc.save('КП.docx') # сохраняю документ
